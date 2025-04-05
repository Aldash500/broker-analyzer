import requests
import pandas as pd
from datetime import datetime
from collections import Counter
import nltk
import re
from docx import Document
import streamlit as st

nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords

st.title("Брокерские новости за 2023 год — Word-отчет")

API_KEY = st.text_input("Введите SerpAPI ключ")

if API_KEY:
    query = "брокер OR брокерская компания OR брокерские услуги"
    months = [
        ("01/01/2023", "01/31/2023", "Январь"),
        ("02/01/2023", "02/28/2023", "Февраль"),
        ("03/01/2023", "03/31/2023", "Март"),
        ("04/01/2023", "04/30/2023", "Апрель")
    ]

    all_data = []

    for start, end, label in months:
        params = {
            "q": query,
            "hl": "ru",
            "tbm": "nws",
            "tbs": f"cdr:1,cd_min:{start},cd_max:{end}",
            "api_key": API_KEY,
            "num": 100
        }
        response = requests.get("https://serpapi.com/search", params=params)
        results = response.json()
        texts = [n.get("snippet", "") for n in results.get("news_results", [])]
        tokens = nltk.word_tokenize(" ".join(texts).lower())
        filtered = [w for w in tokens if w.isalpha() and w not in stopwords.words('russian') and len(w) > 3]
        word_freq = Counter(filtered).most_common(5)
        companies = ["freedom", "финам", "тинькофф", "сбер", "открытие"]
        mentions = {c: sum(t.lower().count(c) for t in texts) for c in companies}
        pos_words = ["прибыль", "рост", "успешно", "инвестиции"]
        neg_words = ["убыток", "снижение", "штраф", "проблема"]
        pos = sum(" ".join(texts).count(w) for w in pos_words)
        neg = sum(" ".join(texts).count(w) for w in neg_words)

        all_data.append({
            "month": label,
            "count": len(texts),
            "pos": pos,
            "neg": neg,
            "words": word_freq,
            "companies": mentions
        })

    # Генерация Word-документа
    doc = Document()
    doc.add_heading("Анализ брокерских новостей за 2023 год", 0)

    for item in all_data:
        doc.add_heading(item["month"], level=1)
        doc.add_paragraph(f"Количество новостей: {item['count']}")
        doc.add_paragraph(f"Позитивные упоминания: {item['pos']}")
        doc.add_paragraph(f"Негативные упоминания: {item['neg']}")
        doc.add_paragraph("Топ-слова: " + ", ".join([f"{w[0]} ({w[1]})" for w in item["words"]]))
        company_line = ", ".join([f"{k}: {v}" for k, v in item["companies"].items() if v > 0]) or "нет"
        doc.add_paragraph("Упомянутые компании: " + company_line)

    word_path = "/mnt/data/broker_report_2023.docx"
    doc.save(word_path)

    st.success("Word-отчет успешно создан!")
    with open(word_path, "rb") as f:
        st.download_button("Скачать Word-отчет", f, file_name="broker_report_2023.docx")
