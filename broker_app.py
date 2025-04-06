
import streamlit as st
from datetime import datetime
from docx import Document
from io import BytesIO

months = {
    "Январь": "01", "Февраль": "02", "Март": "03",
    "Апрель": "04", "Май": "05", "Июнь": "06",
    "Июль": "07", "Август": "08", "Сентябрь": "09",
    "Октябрь": "10", "Ноябрь": "11", "Декабрь": "12"
}

st.title("Генерация Word-отчёта по брокерам")
year = st.selectbox("Выберите год", list(range(2020, datetime.now().year + 1))[::-1])
month_name = st.selectbox("Выберите месяц", list(months.keys()))

if st.button("Сформировать отчёт"):
    doc = Document()
    doc.add_heading(f"Broker Analytics Report — {month_name} {year}", 0)
    doc.add_paragraph(f"Период: {month_name} {year}")
    doc.add_paragraph("Всего новостей: 18")
    doc.add_paragraph("Позитивных упоминаний: 7")
    doc.add_paragraph("Негативных упоминаний: 3")
    doc.add_heading("Примеры позитивных упоминаний", level=2)
    doc.add_paragraph("[EN] Freedom Holding Corp. reported profit growth due to expansion in Kazakhstan.")
    doc.add_paragraph("[RU] Freedom Holding Corp. сообщил о росте прибыли благодаря расширению в Казахстане.")
    doc.add_heading("Источники и аннотации", level=2)
    doc.add_paragraph("Источник: forbes.kz (СМИ)")
    doc.add_paragraph("Аннотация: Freedom расширяет брокерские услуги в регионах. (RU)")
    doc.add_paragraph("Ссылка: https://forbes.kz/news/freedom_expansion")
    doc.add_heading("Выводы и рекомендации", level=2)
    doc.add_paragraph("Рекомендации: Инвестировать в поддержку клиентов, улучшение мобильных решений.")
    doc.add_heading("Полный список источников", level=2)
    doc.add_paragraph("- https://forbes.kz/news/freedom_expansion")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("Скачать Word-отчёт", data=buffer, file_name=f"Broker_Report_{year}_{months[month_name]}.docx")
